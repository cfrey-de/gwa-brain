"""Shared fixtures: an in-memory brain (Qdrant :memory: + lexical embedder + MockLLM)
and document fixtures generated at test time. No key, no network, fully deterministic."""
import pytest

from gwa.embedder import LexicalEmbedder
from gwa.graph.brain import KnowledgeBrain
from gwa.llm import MockLLM


@pytest.fixture
def settings():
    from gwa.config import Settings
    return Settings()


@pytest.fixture
def llm():
    return MockLLM()


@pytest.fixture
def embedder():
    return LexicalEmbedder()


@pytest.fixture
def qdrant():
    from qdrant_client import QdrantClient
    return QdrantClient(location=":memory:")


@pytest.fixture
def brain(tmp_path, qdrant, embedder):
    return KnowledgeBrain(qdrant, embedder, data_dir=str(tmp_path / "data"),
                          collection="test_facts")


@pytest.fixture
def text_doc(tmp_path):
    p = tmp_path / "Bericht.txt"
    p.write_text(
        "Der Wasserstand sinkt nach 500 Stunden auf 160 Zentimeter.\n\n"
        "Der Wasserstand betraegt nach 100 Stunden noch 180 Zentimeter.\n\n"
        "Der Zulauf besteht aus einem verzinkten Stahlrohr.\n\n"
        "Die Ablagerung waechst mit der Betriebsdauer.\n",
        encoding="utf-8",
    )
    return str(p), "Bericht.txt"


@pytest.fixture
def docx_doc(tmp_path):
    import docx
    d = docx.Document()
    d.add_paragraph("Der Betriebsdruck der Pumpe betraegt 3.7 bar.")
    d.add_paragraph("Die Anlage wird bei 25 Grad Celsius getestet.")
    p = tmp_path / "Datenblatt.docx"
    d.save(str(p))
    return str(p), "Datenblatt.docx"


@pytest.fixture
def pdf_doc(tmp_path):
    reportlab = pytest.importorskip("reportlab")  # noqa: F841
    from reportlab.pdfgen import canvas
    p = tmp_path / "Studie.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(72, 720, "Das Wartungsintervall betraegt mindestens 2000 Betriebsstunden.")
    c.drawString(72, 700, "Der Durchfluss liegt unter 5 Litern pro Sekunde.")
    c.showPage()
    c.save()
    return str(p), "Studie.pdf"
